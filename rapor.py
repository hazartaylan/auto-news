import re
import html
import time
import os
import io
import argparse
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone, timedelta
from typing import List, Tuple
from urllib.parse import urljoin

import feedparser
import requests
from bs4 import BeautifulSoup
from PIL import Image

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# cf ai configleri
CF_ACCOUNT_ID = os.getenv("CF_ACCOUNT_ID")
CF_API_TOKEN = os.getenv("CF_API_TOKEN")
CF_MODEL = "@cf/meta/llama-3-8b-instruct"

CF_AI_URL = f"https://api.cloudflare.com/client/v4/accounts/{CF_ACCOUNT_ID}/ai/run/{CF_MODEL}"
CF_HEADERS = {
    "Authorization": f"Bearer {CF_API_TOKEN}",
    "Content-Type": "application/json"
}



@dataclass
class Haber:
    kaynak: str
    baslik: str
    link: str
    tarih: datetime
    ozet: str
    entry: object


# RSS haber kaynakları
RSS_KAYNAKLAR: List[Tuple[str, str]] = [
    ("The Hacker News", "https://thehackernews.com/feeds/posts/default?alt=rss"),
    ("Krebs on Security", "https://krebsonsecurity.com/feed/"),
    ("PortSwigger Research", "https://portswigger.net/research/rss"),
]



def simdi_utc():
    return datetime.now(timezone.utc)


def temiz_metin(s: str) -> str:
    if not s:
        return ""
    s = html.unescape(s)
    s = re.sub(r"\s+", " ", s).strip()
    s = re.sub(r"^(Siber Güvenlik.*?:)", "", s, flags=re.IGNORECASE)
    return s.lstrip("*#- ").strip()


def entry_tarihi(entry):
    for k in ("published_parsed", "updated_parsed"):
        if hasattr(entry, k) and getattr(entry, k):
            return datetime.fromtimestamp(time.mktime(getattr(entry, k)), tz=timezone.utc)
    return None


def fmt_tarih(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


# haber metni çekme
def tam_haber_metni_getir(url: str) -> str:
    try:
        r = requests.get(url, timeout=12, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "lxml")

        for t in soup(["script", "style", "noscript"]):
            t.decompose()

        blocks = []
        for tag in ("article", "main"):
            el = soup.find(tag)
            if el:
                blocks.append(el.get_text(" ", strip=True))

        if not blocks:
            divs = soup.find_all("div")
            longest = max(divs, key=lambda d: len(d.get_text(strip=True)), default=None)
            if longest:
                blocks.append(longest.get_text(" ", strip=True))

        return temiz_metin(max(blocks, key=len, default=""))[:3500]
    except Exception:
        return ""


# ai prompt özet
def ai_ozet_ve_ceviri_cf(haber_metin: str) -> str:
    if not CF_ACCOUNT_ID or not CF_API_TOKEN or len(haber_metin) < 300:
        return ""

    prompt = (
        "Aşağıdaki metni tamamen Türkçe olacak şekilde yeniden yaz ve özetle.\n\n"
        "Haber özeti oluştururken bu kuralları zorunlu kullanacaksın:\n"
        "- Tek paragraf olacak\n"
        "- 4–6 tam cümle olacak\n"
        "- Haber Başlık, haber giriş, haber uyarı, haber duyuru, haber analiz etiketi yazıları yazılmayacak özete eklenmeyecek.\n"
        "- **, ##, markdown, madde işareti haber özetinde sakın kullanma.\n"
        "- Yarım cümle bırakma her cümleyi tam yaz.\n"
        "- Kurumsal siber güvenlik / TI dili kullanarak haber özetini çıkart.\n\n"
        f"Metin:\n{haber_metin}"
    )

    payload = {
        "messages": [
            {"role": "system", "content": "You are a cybersecurity threat intelligence analyst."},
            {"role": "user", "content": prompt}
        ]
    }

    try:
        r = requests.post(CF_AI_URL, headers=CF_HEADERS, json=payload, timeout=45)
        text = temiz_metin(r.json().get("result", {}).get("response", ""))
        if text and text[-1] not in ".!?":
            text = re.sub(r"[.!?][^.!?]*$", ".", text)
        return text
    except Exception:
        return ""


# resim çekme
def gorsel_url_bul(entry, link: str) -> str:
    for attr in ("media_content", "media_thumbnail"):
        if hasattr(entry, attr):
            arr = getattr(entry, attr) or []
            if arr and arr[0].get("url"):
                return arr[0]["url"]

    if hasattr(entry, "links"):
        for l in entry.links:
            if l.get("rel") == "enclosure":
                href = l.get("href")
                if href and href.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
                    return href

    try:
        r = requests.get(link, timeout=8, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "lxml")
        tag = soup.find("meta", property="og:image")
        if tag and tag.get("content"):
            return urljoin(link, tag["content"])
    except Exception:
        pass

    return ""


def gorsel_indir_tmp(url: str) -> str:
    try:
        r = requests.get(url, timeout=8)
        if "image" not in (r.headers.get("Content-Type") or ""):
            return ""

        if len(r.content) > 6 * 1024 * 1024:
            return ""

        img = Image.open(io.BytesIO(r.content)).convert("RGB")
        fd, path = tempfile.mkstemp(suffix=".jpg")
        with os.fdopen(fd, "wb") as f:
            img.save(f, "JPEG", quality=85)
        return path
    except Exception:
        return ""


def karta_gorsel_ekle(right_cell, entry, link: str):
    url = gorsel_url_bul(entry, link)
    if not url:
        return

    path = gorsel_indir_tmp(url)
    if not path:
        return

    try:
        table = right_cell.add_table(1, 1)
        table.autofit = False
        table.columns[0].width = Inches(4.9)
        cell = table.cell(0, 0)

        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for s in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{s}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:color"), "D1D5DB")
            borders.append(b)
        tcPr.append(borders)

        cell.paragraphs[0].add_run().add_picture(path, width=Inches(4.9))
    finally:
        os.remove(path)


# DOCX tasarımı
def set_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for k, v in (("top", 100), ("start", 120), ("bottom", 100), ("end", 120)):
        node = OxmlElement(f"w:{k}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def shade_cell(cell, color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color.replace("#", ""))
    tcPr.append(shd)


# DOCX çıktı
def docx_olustur(items, out_path, gun):
    doc = Document()

    doc.add_heading("Haftalık Siber Güvenlik Haber Özeti", 0)
    sub = doc.add_paragraph(f"Son {gun} gün • {fmt_tarih(simdi_utc())}")
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph("")

    for i, it in enumerate(items, 1):
        table = doc.add_table(1, 2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Inches(1.4)
        table.columns[1].width = Inches(5.8)

        left, right = table.cell(0, 0), table.cell(0, 1)
        shade_cell(left, "#111827")
        shade_cell(right, "#F3F4F6")
        set_cell_margins(left)
        set_cell_margins(right)

        lp = left.paragraphs[0]
        r1 = lp.add_run(it.kaynak.upper())
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(255, 255, 255)

        lp2 = left.add_paragraph(fmt_tarih(it.tarih))
        lp2.runs[0].font.size = Pt(9)
        lp2.runs[0].font.color.rgb = RGBColor(209, 213, 219)

        rp = right.paragraphs[0]
        tr = rp.add_run(f"{i}. {it.baslik}")
        tr.bold = True
        tr.font.size = Pt(12)
        tr.font.color.rgb = RGBColor(17, 24, 39)

        karta_gorsel_ekle(right, it.entry, it.link)

        sp = right.add_paragraph(it.ozet or "—")
        sp.runs[0].font.size = Pt(10)
        sp.runs[0].font.color.rgb = RGBColor(55, 65, 81)

        doc.add_paragraph("")

    doc.save(out_path)


# fetch
def tumunu_cek(gun: int, limit: int):
    items = []
    cutoff = simdi_utc() - timedelta(days=gun)

    for kaynak, feed_url in RSS_KAYNAKLAR:
        feed = feedparser.parse(feed_url)
        for e in feed.entries[:limit]:
            tarih = entry_tarihi(e)
            if not tarih or tarih < cutoff:
                continue

            link = e.get("link", "")
            metin = tam_haber_metni_getir(link)
            ozet = ai_ozet_ve_ceviri_cf(metin) or temiz_metin(e.get("summary", ""))

            items.append(
                Haber(
                    kaynak,
                    temiz_metin(e.get("title", "")),
                    link,
                    tarih,
                    ozet,
                    e
                )
            )
    return items


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--gun", type=int, default=7)
    ap.add_argument("--limit", type=int, default=15)
    ap.add_argument("--out", default="haftalik_siber_haberler.docx")
    args = ap.parse_args()

    items = tumunu_cek(args.gun, args.limit)
    docx_olustur(items, args.out, args.gun)
    print(f"Rapor oluşturuldu: {args.out} ({len(items)} haber)")


if __name__ == "__main__":
    main()
